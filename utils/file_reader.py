from typing import Optional, Union
from io import BytesIO
import PyPDF2
import docx
import csv
import os
import io

def extract_text_from_file(file: Union[str, bytes], filename: Optional[str] = None) -> Optional[str]:
        if isinstance(file, str):
            ext = os.path.splitext(file)[-1].lower()
        elif isinstance(file, bytes) and filename:
            ext = os.path.splitext(filename)[-1].lower()
        else:
            raise ValueError("Invalid input: Provide either a file path or (bytes + filename)")


        if ext == ".txt":
               if isinstance(file, str):
                    with open(file, "r", encoding="utf-8") as f:
                        return f.read()
               else:
                return file.decode("utf-8")

        elif ext == ".docx":
                if isinstance(file, str):
                    doc = docx.Document(file)
                else:
                    doc = docx.Document(BytesIO(file))
                return "\n".join([para.text for para in doc.paragraphs])

        elif ext == ".pdf":
                if isinstance(file, str):
                    with open(file, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                else:
                    reader = PyPDF2.PdfReader(BytesIO(file))
                return "\n".join(page.extract_text() or "" for page in reader.pages)

        else:
                raise ValueError(f"Unsupported file format: {ext}")



def parse_csv_records(file_bytes: bytes):
    records = []
    file_stream = io.StringIO(file_bytes.decode("utf-8-sig"))
    reader = csv.DictReader(file_stream)

    for row in reader:
        record = {
            "Summary": row.get("Summary", "").strip(),
            "Project": row.get("Project", "").strip(),
            "Description": row.get("Description", "").strip(),
            "Sprint": row.get("Sprint", "").strip(),
            "Story point estimate": row.get("Story point estimate", 0),
            "Priority": row.get("Priority", "").strip(),
            "Components": row.get("Components", "").strip(),
            "Fix Versions": row.get("Fix Versions", "").strip(),
            "Assignee": row.get("Assignee", "").strip(),
            "Acceptance Criteria": row.get("Acceptance Criteria", "").strip(),
            "Label": [label.strip() for label in row.get("Label", "").split(",") if label.strip()],
            "Parent": row.get("Parent", "").strip()
        }
        records.append(record)
    return records
